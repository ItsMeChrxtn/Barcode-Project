import csv
import ntplib
import os
import random
import uuid
from datetime import datetime, date, timezone, timedelta
from functools import wraps
from io import StringIO, BytesIO

from bson import ObjectId
from flask import (
    Flask,
    Response,
    flash,
    g,
    redirect,
    render_template,
    request,
    send_from_directory,
    session,
    url_for,
)
import barcode
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask_cors import CORS
from pymongo import MongoClient, ReturnDocument, ASCENDING, DESCENDING
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
BARCODE_DIR = os.path.join(BASE_DIR, "static", "barcodes")
TOOL_IMAGE_DIR = os.path.join(BASE_DIR, "static", "tool_images")
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")
TOOL_CODE_PREFIX = "TL-"
CATEGORY_OPTIONS = [
    "Hand Tools",
    "Power Tools",
    "Electrical",
    "Electronics",
    "Safety",
    "Measuring Tools",
    "Cutting Tools",
    "Fastening Tools",
    "Plumbing",
    "Other",
]

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("FLASK_SECRET_KEY", "replace-this-in-production")
app.config["SESSION_COOKIE_SAMESITE"] = os.environ.get("SESSION_COOKIE_SAMESITE", "Lax")
app.config["SESSION_COOKIE_SECURE"] = os.environ.get("SESSION_COOKIE_SECURE", "0") == "1"


def parse_cors_origins(raw_origins):
    if not raw_origins:
        return []
    return [origin.strip() for origin in raw_origins.split(",") if origin.strip()]


cors_origins = parse_cors_origins(os.environ.get("CORS_ALLOWED_ORIGINS", ""))
if cors_origins:
    CORS(app, resources={r"/api/*": {"origins": cors_origins}}, supports_credentials=True)

# ---------------------------------------------------------------------------
# MongoDB connection
# ---------------------------------------------------------------------------

_mongo_client = None
_indexes_ready = False


def get_mongo_client():
    global _mongo_client
    if _mongo_client is None:
        uri = os.environ.get("MONGODB_URI", "mongodb://localhost:27017")
        _mongo_client = MongoClient(uri)
    return _mongo_client


def get_db():
    if "db" not in g:
        client = get_mongo_client()
        db_name = os.environ.get("MONGODB_DB", "borrow_return_db")
        g.db = client[db_name]
        _ensure_ready(g.db)
    return g.db


def _ensure_ready(db):
    global _indexes_ready
    if _indexes_ready:
        return
    db.admins.create_index("username", unique=True)
    db.tools.create_index("tool_code", unique=True)
    db.tools.create_index("barcode", unique=True)
    db.borrowers.create_index("borrower_id", unique=True)
    if db.admins.find_one({"username": "admin"}) is None:
        db.admins.insert_one({
            "first_name": "System",
            "last_name": "Admin",
            "username": "admin",
            "password_hash": generate_password_hash("admin123"),
        })
    _indexes_ready = True


@app.teardown_appcontext
def close_db(error=None):
    _ = error
    g.pop("db", None)


# ---------------------------------------------------------------------------
# MongoDB document helpers
# ---------------------------------------------------------------------------

def to_doc(d):
    if d is None:
        return None
    d = dict(d)
    if "_id" in d:
        d["id"] = str(d.pop("_id"))
    return d


def to_docs(cursor):
    return [to_doc(d) for d in cursor]


def to_oid(s):
    try:
        return ObjectId(str(s))
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Auth helpers
# ---------------------------------------------------------------------------

def login_required(view_func):
    @wraps(view_func)
    def wrapped_view(*args, **kwargs):
        if not session.get("admin_id"):
            flash("Please log in to continue.", "warning")
            return redirect(url_for("login"))
        return view_func(*args, **kwargs)
    return wrapped_view


# ---------------------------------------------------------------------------
# Tool helpers
# ---------------------------------------------------------------------------

def normalize_tool_status(tool):
    return "Available" if tool.get("available_quantity", 0) > 0 else "Unavailable"


def update_tool_status(db, tool_id):
    oid = to_oid(tool_id)
    if not oid:
        return
    tool = db.tools.find_one({"_id": oid})
    if not tool:
        return
    available_qty = max(0, min(tool["available_quantity"], tool["quantity"]))
    status = "Available" if available_qty > 0 else "Unavailable"
    db.tools.update_one(
        {"_id": oid},
        {"$set": {"available_quantity": available_qty, "status": status}},
    )


def get_next_tool_number(db):
    pipeline = [
        {"$match": {"tool_code": {"$regex": f"^{TOOL_CODE_PREFIX}\\d+$"}}},
        {"$project": {"num": {"$toInt": {"$substrCP": ["$tool_code", len(TOOL_CODE_PREFIX), 10]}}}},
        {"$group": {"_id": None, "max_num": {"$max": "$num"}}},
    ]
    result = list(db.tools.aggregate(pipeline))
    if result and result[0].get("max_num") is not None:
        return result[0]["max_num"] + 1
    return 0


def format_tool_code(number):
    return f"{TOOL_CODE_PREFIX}{number:03d}"


def generate_next_tool_code(db):
    return format_tool_code(get_next_tool_number(db))


def get_category_options(selected_category=None):
    options = list(CATEGORY_OPTIONS)
    if selected_category and selected_category not in options:
        options.append(selected_category)
    return options


def generate_unique_barcode(db, length=12):
    while True:
        value = "".join(str(random.randint(0, 9)) for _ in range(length))
        if db.tools.find_one({"barcode": value}) is None:
            return value


def create_barcode_image(barcode_value):
    os.makedirs(BARCODE_DIR, exist_ok=True)
    file_stem = f"barcode_{barcode_value}"
    path_without_ext = os.path.join(BARCODE_DIR, file_stem)
    code128 = barcode.get("code128", barcode_value, writer=ImageWriter())
    generated_path = code128.save(path_without_ext)
    return os.path.basename(generated_path)


def save_tool_image(file_storage):
    if file_storage is None or not getattr(file_storage, "filename", ""):
        return None
    filename = secure_filename(file_storage.filename)
    if not filename or "." not in filename:
        return None
    extension = filename.rsplit(".", 1)[1].lower()
    if extension not in {"png", "jpg", "jpeg", "webp"}:
        return None
    os.makedirs(TOOL_IMAGE_DIR, exist_ok=True)
    generated_name = f"tool_{uuid.uuid4().hex[:12]}.{extension}"
    save_path = os.path.join(TOOL_IMAGE_DIR, generated_name)
    file_storage.save(save_path)
    return generated_name


def get_tool_by_barcode(db, barcode_value):
    return to_doc(db.tools.find_one({"barcode": barcode_value.strip()}))


def build_tool_payload(tool):
    return {
        "id": tool["id"],
        "tool_name": tool["tool_name"],
        "tool_code": tool["tool_code"],
        "category": tool["category"],
        "quantity": tool["quantity"],
        "available_quantity": tool["available_quantity"],
        "barcode": tool["barcode"],
        "barcode_image": tool.get("barcode_image"),
        "tool_image": tool.get("tool_image"),
        "status": normalize_tool_status(tool),
    }


def build_tool_filters(search, category, availability):
    filters = {}
    if search:
        regex = {"$regex": search, "$options": "i"}
        filters["$or"] = [
            {"tool_name": regex},
            {"tool_code": regex},
            {"barcode": regex},
            {"category": regex},
            {"status": regex},
        ]
    if category:
        filters["category"] = category
    if availability == "available":
        filters["available_quantity"] = {"$gt": 0}
    elif availability == "unavailable":
        filters["available_quantity"] = {"$lte": 0}
    return filters


def parse_tool_ids(raw_ids):
    parsed = []
    seen = set()
    for raw in raw_ids:
        s = str(raw).strip()
        if not s or s in seen:
            continue
        if to_oid(s) is None:
            continue
        parsed.append(s)
        seen.add(s)
    return parsed


def split_deletable_tool_ids(db, tool_ids):
    if not tool_ids:
        return [], []
    oids = [to_oid(tid) for tid in tool_ids if to_oid(tid)]
    rows = list(db.tools.find({"_id": {"$in": oids}}))
    borrowed_ids = set()
    for row in rows:
        if db.transactions.count_documents({"tool_id": row["_id"], "status": "borrowed"}) > 0:
            borrowed_ids.add(str(row["_id"]))
    deletable_ids = [tid for tid in tool_ids if tid not in borrowed_ids]
    blocked_tools = [to_doc(r) for r in rows if str(r["_id"]) in borrowed_ids]
    return deletable_ids, blocked_tools


# ---------------------------------------------------------------------------
# Transaction join pipeline helpers
# ---------------------------------------------------------------------------

def _transaction_join_stages():
    return [
        {"$lookup": {"from": "borrowers", "localField": "borrower_id", "foreignField": "_id", "as": "borrower"}},
        {"$lookup": {"from": "tools", "localField": "tool_id", "foreignField": "_id", "as": "tool_doc"}},
        {"$lookup": {"from": "admins", "localField": "lent_by_admin_id", "foreignField": "_id", "as": "admin_doc"}},
        {"$unwind": {"path": "$borrower", "preserveNullAndEmptyArrays": True}},
        {"$unwind": {"path": "$tool_doc", "preserveNullAndEmptyArrays": True}},
        {"$unwind": {"path": "$admin_doc", "preserveNullAndEmptyArrays": True}},
        {"$project": {
            "borrow_date": 1,
            "expected_return_date": 1,
            "return_date": 1,
            "status": 1,
            "barcode": 1,
            "borrower_name": "$borrower.borrower_name",
            "borrower_id": "$borrower.borrower_id",
            "course_department": "$borrower.course_department",
            "contact_number": "$borrower.contact_number",
            "tool_id": {"$toString": "$tool_doc._id"},
            "tool_name": "$tool_doc.tool_name",
            "tool_code": "$tool_doc.tool_code",
            "category": "$tool_doc.category",
            "lent_by_admin_name": {
                "$trim": {
                    "input": {
                        "$concat": [
                            {"$ifNull": ["$admin_doc.first_name", ""]},
                            " ",
                            {"$ifNull": ["$admin_doc.last_name", ""]},
                        ]
                    }
                }
            },
        }},
    ]


def get_active_borrow_transaction(db, barcode_value):
    pipeline = [
        {"$match": {"barcode": barcode_value.strip(), "status": "borrowed"}},
        {"$sort": {"borrow_date": ASCENDING, "_id": ASCENDING}},
        {"$limit": 1},
        *_transaction_join_stages(),
    ]
    result = list(db.transactions.aggregate(pipeline))
    return to_doc(result[0]) if result else None


def get_transaction_with_joins(db, transaction_oid):
    pipeline = [
        {"$match": {"_id": transaction_oid}},
        *_transaction_join_stages(),
    ]
    result = list(db.transactions.aggregate(pipeline))
    return to_doc(result[0]) if result else None


def build_transaction_payload(transaction):
    return {
        "id": transaction["id"],
        "borrow_date": transaction.get("borrow_date"),
        "expected_return_date": transaction.get("expected_return_date"),
        "return_date": transaction.get("return_date"),
        "status": transaction.get("status"),
        "borrower_name": transaction.get("borrower_name"),
        "borrower_id": transaction.get("borrower_id"),
        "course_department": transaction.get("course_department"),
        "contact_number": transaction.get("contact_number"),
        "tool_id": transaction.get("tool_id"),
        "tool_name": transaction.get("tool_name"),
        "tool_code": transaction.get("tool_code"),
        "category": transaction.get("category"),
        "barcode": transaction.get("barcode"),
        "lent_by_admin_name": transaction.get("lent_by_admin_name"),
    }


# ---------------------------------------------------------------------------
# Borrower helpers
# ---------------------------------------------------------------------------

def upsert_borrower(db, borrower_name, borrower_code, course_department, contact_number):
    result = db.borrowers.find_one_and_update(
        {"borrower_id": borrower_code},
        {"$set": {
            "borrower_name": borrower_name,
            "borrower_id": borrower_code,
            "course_department": course_department,
            "contact_number": contact_number,
        }},
        upsert=True,
        return_document=ReturnDocument.AFTER,
    )
    return result["_id"]


# ---------------------------------------------------------------------------
# Borrow / Return scan processors
# ---------------------------------------------------------------------------

def process_borrow_scan(db, payload):
    barcode_value = payload.get("barcode", "").strip()
    borrower_name = payload.get("borrower_name", "").strip()
    borrower_code = payload.get("borrower_id", "").strip()
    course_department = payload.get("course_department", "").strip()
    contact_number = payload.get("contact_number", "").strip()

    if not all([borrower_name, borrower_code, course_department, contact_number]):
        return {"ok": False, "message": "Please complete the borrower information before scanning.", "category": "warning"}, 400

    tool = get_tool_by_barcode(db, barcode_value)
    if tool is None:
        return {"ok": False, "message": "No tool found for the scanned barcode.", "category": "error"}, 404

    active_transaction = get_active_borrow_transaction(db, barcode_value)
    if active_transaction is not None:
        return {
            "ok": False,
            "message": "This barcode is already borrowed.",
            "category": "warning",
            "tool": build_tool_payload(tool),
            "transaction": build_transaction_payload(active_transaction),
        }, 409

    if tool["available_quantity"] <= 0:
        return {
            "ok": False,
            "message": "This tool is currently unavailable.",
            "category": "warning",
            "tool": build_tool_payload(tool),
        }, 409

    borrower_oid = upsert_borrower(db, borrower_name, borrower_code, course_department, contact_number)
    borrow_timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    admin_id_str = session.get("admin_id")
    lending_admin_oid = to_oid(admin_id_str) if admin_id_str else None

    result = db.transactions.insert_one({
        "borrower_id": borrower_oid,
        "tool_id": to_oid(tool["id"]),
        "barcode": barcode_value,
        "borrow_date": borrow_timestamp,
        "expected_return_date": "",
        "return_date": None,
        "status": "borrowed",
        "lent_by_admin_id": lending_admin_oid,
    })
    db.tools.update_one(
        {"_id": to_oid(tool["id"])},
        {"$inc": {"available_quantity": -1}},
    )
    update_tool_status(db, tool["id"])

    updated_tool = get_tool_by_barcode(db, barcode_value)
    saved_transaction = get_transaction_with_joins(db, result.inserted_id)

    return {
        "ok": True,
        "action": "borrowed",
        "message": "Borrow record saved successfully.",
        "category": "success",
        "tool": build_tool_payload(updated_tool),
        "transaction": build_transaction_payload(saved_transaction),
    }, 200


def process_return_scan(db, payload):
    barcode_value = payload.get("barcode", "").strip()
    tool = get_tool_by_barcode(db, barcode_value)

    if tool is None:
        return {"ok": False, "message": "No tool found for the scanned barcode.", "category": "error"}, 404

    active_transaction = get_active_borrow_transaction(db, barcode_value)
    if active_transaction is None:
        return {
            "ok": False,
            "message": "This barcode is not currently borrowed.",
            "category": "warning",
            "tool": build_tool_payload(tool),
        }, 409

    returned_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    transaction_oid = to_oid(active_transaction["id"])
    db.transactions.update_one(
        {"_id": transaction_oid},
        {"$set": {"return_date": returned_at, "status": "returned"}},
    )
    db.tools.update_one(
        {"_id": to_oid(active_transaction["tool_id"])},
        {"$inc": {"available_quantity": 1}},
    )
    update_tool_status(db, active_transaction["tool_id"])

    updated_tool = get_tool_by_barcode(db, barcode_value)
    saved_transaction = get_transaction_with_joins(db, transaction_oid)

    return {
        "ok": True,
        "action": "returned",
        "message": "Return record saved successfully.",
        "category": "success",
        "tool": build_tool_payload(updated_tool),
        "transaction": build_transaction_payload(saved_transaction),
    }, 200


# ---------------------------------------------------------------------------
# Page render helpers
# ---------------------------------------------------------------------------

def render_add_tool_page(db, generated_tool=None, form_data=None):
    return render_template(
        "add_tool.html",
        generated_tool=generated_tool,
        next_tool_code=generate_next_tool_code(db),
        category_options=get_category_options((form_data or {}).get("category")),
        form_data=form_data or {},
    )


def render_admins_page(db, form_data=None):
    admins = to_docs(db.admins.find({}, {"username": 1, "first_name": 1, "last_name": 1}).sort("username", ASCENDING))
    return render_template(
        "admins.html",
        admins=admins,
        form_data=form_data or {},
    )


# ---------------------------------------------------------------------------
# Misc helpers
# ---------------------------------------------------------------------------

def parse_date_input(value):
    if not value:
        return None
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except ValueError:
        return None


def get_ntp_now():
    try:
        c = ntplib.NTPClient()
        response = c.request("pool.ntp.org", version=3, timeout=2)
        return datetime.fromtimestamp(response.tx_time)
    except Exception:
        return datetime.now()


def get_project_profiles():
    return [
        {"name": "BERTOLDO, Restygie D.", "role": "Proponent", "image": "BERTOLDO, Restygie D..jpg"},
        {"name": "CASTUERAS, Rhenard R.", "role": "Proponent", "image": "CASTUERAS, Rhenard R..jpg"},
        {"name": "VEDAD, Kim Jay C.", "role": "Proponent", "image": "VEDAD, Kim Jay C..jpg"},
        {"name": "Nicky Jay R. Evangelista", "role": "Adviser", "image": "Nicky Jay R. Evangelista - Adviser.jpg"},
        {"name": "Diane P. Arayata", "role": "Technical Critic", "image": "Diane P. Arayata - Technical critic.jpg"},
    ]


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    return render_template("landing.html", profiles=get_project_profiles())


@app.route("/title-page")
def title_page():
    return render_template("landing.html", profiles=get_project_profiles())


@app.route("/template-image/<path:filename>")
def template_image(filename):
    return send_from_directory(TEMPLATE_DIR, filename)


@app.route("/healthz")
def healthz():
    return {"ok": True, "service": "borrow-return-system"}, 200


@app.route("/login", methods=["GET", "POST"])
def login():
    if session.get("admin_id"):
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")

        if not username or not password:
            flash("Username and password are required.", "danger")
            return render_template("login.html")

        db = get_db()
        admin = db.admins.find_one({"username": username})

        if admin is None or not check_password_hash(admin["password_hash"], password):
            flash("Invalid username or password.", "danger")
            return render_template("login.html")

        session.clear()
        session["admin_id"] = str(admin["_id"])
        session["admin_username"] = admin["username"]
        session["admin_full_name"] = f"{admin.get('first_name', '')} {admin.get('last_name', '')}".strip()
        flash("Welcome back!", "success")
        return redirect(url_for("dashboard"))

    return render_template("login.html")


@app.route("/logout")
@login_required
def logout():
    session.clear()
    flash("You have been logged out.", "info")
    return redirect(url_for("login"))


@app.route("/admins", methods=["GET", "POST"])
@login_required
def admins():
    db = get_db()

    if request.method == "POST":
        first_name = request.form.get("first_name", "").strip()
        last_name = request.form.get("last_name", "").strip()
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        confirm_password = request.form.get("confirm_password", "")
        form_data = {"first_name": first_name, "last_name": last_name, "username": username}

        if not first_name or not last_name or not username or not password or not confirm_password:
            flash("First name, last name, username, password, and confirmation are required.", "danger")
            return render_admins_page(db, form_data=form_data)

        if len(first_name) < 2 or len(last_name) < 2:
            flash("First name and last name must be at least 2 characters long.", "danger")
            return render_admins_page(db, form_data=form_data)

        if len(username) < 3:
            flash("Username must be at least 3 characters long.", "danger")
            return render_admins_page(db, form_data=form_data)

        if len(password) < 6:
            flash("Password must be at least 6 characters long.", "danger")
            return render_admins_page(db, form_data=form_data)

        if password != confirm_password:
            flash("Password confirmation does not match.", "danger")
            return render_admins_page(db, form_data=form_data)

        if db.admins.find_one({"username": username}):
            flash("That username is already in use.", "danger")
            return render_admins_page(db, form_data=form_data)

        db.admins.insert_one({
            "first_name": first_name,
            "last_name": last_name,
            "username": username,
            "password_hash": generate_password_hash(password),
        })
        flash("New admin account created successfully.", "success")
        return redirect(url_for("admins"))

    return render_admins_page(db)


@app.route("/admins/<admin_id>/edit", methods=["GET", "POST"])
@login_required
def edit_admin(admin_id):
    current_admin_id = session.get("admin_id")
    if current_admin_id != admin_id:
        flash("You can only edit your own account details.", "warning")
        return redirect(url_for("edit_admin", admin_id=current_admin_id))

    db = get_db()
    admin_oid = to_oid(admin_id)
    admin = to_doc(db.admins.find_one({"_id": admin_oid}, {"username": 1, "first_name": 1, "last_name": 1}))

    if admin is None:
        flash("Admin account not found.", "danger")
        return redirect(url_for("admins"))

    if request.method == "POST":
        first_name = request.form.get("first_name", "").strip()
        last_name = request.form.get("last_name", "").strip()
        username = request.form.get("username", "").strip()
        new_password = request.form.get("new_password", "")
        confirm_password = request.form.get("confirm_password", "")
        form_admin = {"id": admin_id, "first_name": first_name, "last_name": last_name, "username": username}

        if not first_name or not last_name or not username:
            flash("First name, last name, and username are required.", "danger")
            return render_template("admin_edit.html", admin=form_admin)

        if len(first_name) < 2 or len(last_name) < 2:
            flash("First name and last name must be at least 2 characters long.", "danger")
            return render_template("admin_edit.html", admin=form_admin)

        if len(username) < 3:
            flash("Username must be at least 3 characters long.", "danger")
            return render_template("admin_edit.html", admin=form_admin)

        if db.admins.find_one({"username": username, "_id": {"$ne": admin_oid}}):
            flash("That username is already in use.", "danger")
            return render_template("admin_edit.html", admin=form_admin)

        update_fields = {"first_name": first_name, "last_name": last_name, "username": username}

        if new_password or confirm_password:
            if len(new_password) < 6:
                flash("New password must be at least 6 characters long.", "danger")
                return render_template("admin_edit.html", admin=form_admin)
            if new_password != confirm_password:
                flash("Password confirmation does not match.", "danger")
                return render_template("admin_edit.html", admin=form_admin)
            update_fields["password_hash"] = generate_password_hash(new_password)

        db.admins.update_one({"_id": admin_oid}, {"$set": update_fields})
        session["admin_username"] = username
        session["admin_full_name"] = f"{first_name} {last_name}".strip()
        flash("Your account details were updated successfully.", "success")
        return redirect(url_for("admins"))

    return render_template("admin_edit.html", admin=admin)


@app.route("/dashboard")
@login_required
def dashboard():
    db = get_db()

    total_tools = db.tools.count_documents({})
    available_tools = db.tools.count_documents({"available_quantity": {"$gt": 0}})
    borrowed_tools = db.transactions.count_documents({"status": "borrowed"})
    returned_today = db.transactions.count_documents({
        "return_date": {"$regex": f"^{date.today().isoformat()}"},
    })
    total_transactions = db.transactions.count_documents({})
    total_admins = db.admins.count_documents({})

    pipeline = [
        {"$sort": {"_id": DESCENDING}},
        {"$limit": 10},
        {"$lookup": {"from": "borrowers", "localField": "borrower_id", "foreignField": "_id", "as": "borrower"}},
        {"$lookup": {"from": "tools", "localField": "tool_id", "foreignField": "_id", "as": "tool_doc"}},
        {"$unwind": {"path": "$borrower", "preserveNullAndEmptyArrays": True}},
        {"$unwind": {"path": "$tool_doc", "preserveNullAndEmptyArrays": True}},
        {"$project": {
            "borrower_name": "$borrower.borrower_name",
            "borrower_id": "$borrower.borrower_id",
            "tool_name": "$tool_doc.tool_name",
            "barcode": 1,
            "borrow_date": 1,
            "expected_return_date": 1,
            "return_date": 1,
            "status": 1,
        }},
    ]
    recent_transactions = to_docs(db.transactions.aggregate(pipeline))

    return render_template(
        "dashboard.html",
        total_tools=total_tools,
        available_tools=available_tools,
        borrowed_tools=borrowed_tools,
        returned_today=returned_today,
        total_transactions=total_transactions,
        total_admins=total_admins,
        recent_transactions=recent_transactions,
    )


@app.route("/tools")
@login_required
def tools():
    db = get_db()
    search = request.args.get("search", "").strip()
    category = request.args.get("category", "").strip()
    availability = request.args.get("availability", "").strip()

    try:
        page = max(1, int(request.args.get("page", "1")))
    except ValueError:
        page = 1

    per_page = 9
    filters = build_tool_filters(search, category, availability)
    total_items = db.tools.count_documents(filters)
    total_pages = max(1, (total_items + per_page - 1) // per_page)
    page = min(page, total_pages)
    offset = (page - 1) * per_page

    tool_rows = to_docs(db.tools.find(filters).sort("_id", DESCENDING).skip(offset).limit(per_page))
    categories = [c["_id"] for c in db.tools.aggregate([
        {"$group": {"_id": "$category"}},
        {"$sort": {"_id": ASCENDING}},
    ])]

    return render_template(
        "tools.html",
        tools=tool_rows,
        categories=categories,
        category_options=CATEGORY_OPTIONS,
        search=search,
        selected_category=category,
        availability=availability,
        next_tool_code=generate_next_tool_code(db),
        page=page,
        per_page=per_page,
        total_items=total_items,
        total_pages=total_pages,
    )


@app.route("/tools/barcodes/print-selected", methods=["POST"])
@login_required
def print_selected_barcodes():
    db = get_db()
    valid_ids = parse_tool_ids(request.form.getlist("tool_ids"))

    if not valid_ids:
        flash("Select at least one tool barcode to print.", "warning")
        return redirect(url_for("tools"))

    oids = [to_oid(tid) for tid in valid_ids if to_oid(tid)]
    rows = to_docs(db.tools.find({"_id": {"$in": oids}}).sort("_id", DESCENDING))
    return render_template("print_barcodes.html", tools=rows, print_title="Selected Tool Barcodes")


@app.route("/tools/barcodes/print-all")
@login_required
def print_all_barcodes():
    db = get_db()
    search = request.args.get("search", "").strip()
    category = request.args.get("category", "").strip()
    availability = request.args.get("availability", "").strip()
    filters = build_tool_filters(search, category, availability)

    rows = to_docs(db.tools.find(filters).sort("_id", DESCENDING))
    if not rows:
        flash("No tool barcodes available to print for the current filter.", "warning")
        return redirect(url_for("tools", search=search, category=category, availability=availability))

    return render_template("print_barcodes.html", tools=rows, print_title="All Tool Barcodes")


def create_barcodes_word_document(tools, title):
    """Generate a Word document with resizable barcodes"""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"{len(tools)} barcode(s) ready for printing/resizing")
    doc.add_paragraph()  # Add blank line

    # Create a table with 3 columns for better layout
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Set table width and cell height
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(2.0)

    for i, tool in enumerate(tools):
        # Add new row if needed
        if i % 3 == 0 and i != 0:
            row = table.add_row()
        else:
            row = table.rows[-1]

        cell = row.cells[i % 3]
        
        # Clear all paragraphs first
        for _ in range(len(cell.paragraphs)):
            p = cell.paragraphs[0]._element
            p.getparent().remove(p)
        
        # Add barcode image to cell
        if tool.get('barcode_image'):
            barcode_path = os.path.join(BARCODE_DIR, tool['barcode_image'])
            if os.path.exists(barcode_path):
                try:
                    paragraph = cell.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(barcode_path, width=Inches(1.8))
                except Exception as e:
                    error_para = cell.add_paragraph()
                    error_para.text = f"Error: {str(e)}"
        else:
            # If no barcode image, add placeholder
            para = cell.add_paragraph()
            para.text = "[No barcode image]"
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add tool information in a new paragraph
        info_para = cell.add_paragraph()
        info_para.paragraph_format.space_before = Pt(6)
        info_run = info_para.add_run(f"{tool.get('tool_name', 'N/A')}\n")
        info_run.font.size = Pt(9)
        info_run.font.bold = True
        
        info_run2 = info_para.add_run(f"Code: {tool.get('tool_code', 'N/A')}\n")
        info_run2.font.size = Pt(8)
        
        info_run3 = info_para.add_run(f"Barcode: {tool.get('barcode', 'N/A')}")
        info_run3.font.size = Pt(8)

    return doc


@app.route("/tools/barcodes/download-selected-word", methods=["POST"])
@login_required
def export_selected_barcodes_to_word():
    db = get_db()
    valid_ids = parse_tool_ids(request.form.getlist("tool_ids"))

    if not valid_ids:
        flash("Select at least one tool barcode to export.", "warning")
        return redirect(url_for("tools"))

    oids = [to_oid(tid) for tid in valid_ids if to_oid(tid)]
    rows = to_docs(db.tools.find({"_id": {"$in": oids}}).sort("_id", DESCENDING))
    
    doc = create_barcodes_word_document(rows, "Selected Tool Barcodes")
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return Response(
        doc_io.getvalue(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=selected_barcodes.docx"}
    )


@app.route("/tools/barcodes/download-all-word")
@login_required
def export_all_barcodes_to_word():
    db = get_db()
    search = request.args.get("search", "").strip()
    category = request.args.get("category", "").strip()
    availability = request.args.get("availability", "").strip()
    filters = build_tool_filters(search, category, availability)

    rows = to_docs(db.tools.find(filters).sort("_id", DESCENDING))
    if not rows:
        flash("No tool barcodes available to export for the current filter.", "warning")
        return redirect(url_for("tools", search=search, category=category, availability=availability))

    doc = create_barcodes_word_document(rows, "All Tool Barcodes")
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return Response(
        doc_io.getvalue(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=all_barcodes.docx"}
    )


@app.route("/tools/add", methods=["GET", "POST"])
@login_required
def add_tool():
    db = get_db()

    if request.method == "POST":
        tool_name = request.form.get("tool_name", "").strip()
        category = request.form.get("category", "").strip()
        description = request.form.get("description", "").strip()
        quantity_text = request.form.get("quantity", "0").strip()
        uploaded_tool_image = save_tool_image(request.files.get("tool_image"))

        if not all([tool_name, category]):
            flash("Tool name and category are required.", "danger")
            return redirect(url_for("tools"))

        if category not in CATEGORY_OPTIONS:
            flash("Please select a valid category.", "danger")
            return redirect(url_for("tools"))

        try:
            quantity = int(quantity_text)
            if quantity < 1:
                raise ValueError
        except ValueError:
            flash("Quantity must be at least 1.", "danger")
            return redirect(url_for("tools"))

        start_number = get_next_tool_number(db)
        created_codes = []
        created_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        for offset in range(quantity):
            tool_code = format_tool_code(start_number + offset)
            barcode_val = generate_unique_barcode(db)
            barcode_image = create_barcode_image(barcode_val)
            db.tools.insert_one({
                "tool_name": tool_name,
                "tool_code": tool_code,
                "category": category,
                "description": description,
                "quantity": 1,
                "available_quantity": 1,
                "barcode": barcode_val,
                "barcode_image": barcode_image,
                "tool_image": uploaded_tool_image,
                "status": "Available",
                "date_added": created_time,
            })
            created_codes.append(tool_code)

        flash(f"Created {quantity} tool unit(s): {created_codes[0]} to {created_codes[-1]}.", "success")
        return redirect(url_for("tools"))

    return redirect(url_for("tools"))


@app.route("/tools/<tool_id>/edit", methods=["GET", "POST"])
@login_required
def edit_tool(tool_id):
    db = get_db()
    tool_oid = to_oid(tool_id)
    tool = to_doc(db.tools.find_one({"_id": tool_oid}))

    if tool is None:
        flash("Tool not found.", "danger")
        return redirect(url_for("tools"))

    if request.method == "POST":
        tool_name = request.form.get("tool_name", "").strip()
        category = request.form.get("category", "").strip()
        description = request.form.get("description", "").strip()
        quantity_text = request.form.get("quantity", "0").strip()
        barcode_val = request.form.get("barcode", "").strip()
        tool_code = tool["tool_code"]
        form_tool = dict(tool)
        form_tool.update({"tool_name": tool_name, "category": category, "description": description, "quantity": quantity_text, "barcode": barcode_val})

        if not all([tool_name, category, barcode_val]):
            flash("Tool name, category, and barcode are required.", "danger")
            return render_template("edit_tool.html", tool=form_tool, category_options=get_category_options(category))

        if category not in get_category_options(tool["category"]):
            flash("Please select a valid category.", "danger")
            return render_template("edit_tool.html", tool=form_tool, category_options=get_category_options(category))

        try:
            quantity = int(quantity_text)
            if quantity < 0:
                raise ValueError
        except ValueError:
            flash("Quantity must be a non-negative number.", "danger")
            return render_template("edit_tool.html", tool=form_tool, category_options=get_category_options(category))

        if db.tools.find_one({"barcode": barcode_val, "_id": {"$ne": tool_oid}}):
            flash("Barcode already exists. Please use a unique barcode.", "danger")
            return render_template("edit_tool.html", tool=form_tool, category_options=get_category_options(category))

        borrowed_count = db.transactions.count_documents({"tool_id": tool_oid, "status": "borrowed"})
        if quantity < borrowed_count:
            flash("Quantity cannot be less than currently borrowed units.", "danger")
            return render_template("edit_tool.html", tool=form_tool, category_options=get_category_options(category))

        available_quantity = max(0, quantity - borrowed_count)
        status = "Available" if available_quantity > 0 else "Unavailable"
        barcode_image = tool.get("barcode_image")
        if tool["barcode"] != barcode_val or not barcode_image:
            barcode_image = create_barcode_image(barcode_val)

        db.tools.update_one(
            {"_id": tool_oid},
            {"$set": {
                "tool_name": tool_name,
                "tool_code": tool_code,
                "category": category,
                "description": description,
                "quantity": quantity,
                "available_quantity": available_quantity,
                "barcode": barcode_val,
                "barcode_image": barcode_image,
                "status": status,
            }},
        )
        flash("Tool updated successfully.", "success")
        return redirect(url_for("tools"))

    return render_template("edit_tool.html", tool=tool, category_options=get_category_options(tool["category"]))


@app.route("/tools/<tool_id>/delete", methods=["POST"])
@login_required
def delete_tool(tool_id):
    _ = tool_id
    flash("Tool deletion is disabled to preserve inventory history.", "warning")
    return redirect(url_for("tools"))


@app.route("/tools/delete-selected", methods=["POST"])
@login_required
def delete_selected_tools():
    flash("Bulk tool deletion is disabled to preserve inventory history.", "warning")
    return redirect(url_for("tools"))


@app.route("/tools/<tool_id>/barcode/regenerate", methods=["POST"])
@login_required
def regenerate_tool_barcode(tool_id):
    db = get_db()
    tool_oid = to_oid(tool_id)
    tool = db.tools.find_one({"_id": tool_oid}, {"barcode": 1})
    if tool is None:
        flash("Tool not found.", "danger")
        return redirect(url_for("tools"))

    barcode_image = create_barcode_image(tool["barcode"])
    db.tools.update_one({"_id": tool_oid}, {"$set": {"barcode_image": barcode_image}})
    flash("Barcode image generated successfully.", "success")
    return redirect(url_for("tools"))


@app.route("/tools/regenerate-all-barcodes", methods=["GET", "POST"])
@login_required
def regenerate_all_barcodes():
    db = get_db()
    tools = db.tools.find()
    count = 0
    for tool in tools:
        if not tool.get("barcode_image"):
            barcode_image = create_barcode_image(tool["barcode"])
            db.tools.update_one({"_id": tool["_id"]}, {"$set": {"barcode_image": barcode_image}})
            count += 1
    if count > 0:
        flash(f"Regenerated barcode images for {count} tool(s).", "success")
    else:
        flash("All tools already have barcode images.", "info")
    return redirect(url_for("tools"))


@app.route("/borrow", methods=["GET"])
@login_required
def borrow_tool():
    return render_template("borrow.html")


@app.route("/return", methods=["GET"])
@login_required
def return_tool():
    return redirect(url_for("borrow_tool"))


@app.route("/api/scan/process", methods=["POST"])
@login_required
def api_process_scan():
    payload = request.get_json(silent=True) or request.form.to_dict()
    mode = payload.get("mode", "borrow").strip().lower()
    barcode_value = payload.get("barcode", "").strip()

    if mode not in {"borrow", "return"}:
        return {"ok": False, "message": "Invalid scan mode.", "category": "error"}, 400

    if not barcode_value:
        return {"ok": False, "message": "Please scan a barcode first.", "category": "warning"}, 400

    db = get_db()
    if mode == "borrow":
        return process_borrow_scan(db, payload)
    return process_return_scan(db, payload)


@app.route("/transactions")
@login_required
def transactions():
    db = get_db()
    search = request.args.get("search", "").strip()
    status = request.args.get("status", "").strip()
    borrower = request.args.get("borrower", "").strip()
    tool = request.args.get("tool", "").strip()
    date_from = request.args.get("date_from", "").strip()
    date_to = request.args.get("date_to", "").strip()

    pipeline = [*_transaction_join_stages()]

    match = {}
    if search:
        regex = {"$regex": search, "$options": "i"}
        match["$or"] = [
            {"borrower_name": regex},
            {"borrower_id": regex},
            {"tool_name": regex},
            {"barcode": regex},
            {"status": regex},
            {"lent_by_admin_name": regex},
        ]
    if status:
        match["status"] = status
    if borrower:
        match["borrower_id"] = {"$regex": borrower, "$options": "i"}
    if tool:
        match["tool_name"] = {"$regex": tool, "$options": "i"}
    if date_from:
        match.setdefault("borrow_date", {})["$gte"] = date_from
    if date_to:
        match.setdefault("borrow_date", {})["$lte"] = date_to + " 23:59:59"

    if match:
        pipeline.append({"$match": match})
    pipeline.append({"$sort": {"_id": DESCENDING}})

    rows = to_docs(db.transactions.aggregate(pipeline))

    return render_template(
        "transactions.html",
        transactions=rows,
        search=search,
        selected_status=status,
        borrower=borrower,
        tool=tool,
        date_from=date_from,
        date_to=date_to,
    )


@app.route("/reports")
@login_required
def reports():
    db = get_db()

    most_borrowed_pipeline = [
        {"$group": {"_id": "$tool_id", "borrow_count": {"$sum": 1}}},
        {"$sort": {"borrow_count": DESCENDING}},
        {"$limit": 10},
        {"$lookup": {"from": "tools", "localField": "_id", "foreignField": "_id", "as": "tool_doc"}},
        {"$unwind": {"path": "$tool_doc", "preserveNullAndEmptyArrays": True}},
        {"$project": {
            "tool_name": "$tool_doc.tool_name",
            "tool_code": "$tool_doc.tool_code",
            "barcode": "$tool_doc.barcode",
            "borrow_count": 1,
        }},
    ]
    most_borrowed_tools = to_docs(db.transactions.aggregate(most_borrowed_pipeline))

    currently_borrowed_pipeline = [
        {"$match": {"status": "borrowed"}},
        {"$sort": {"borrow_date": ASCENDING}},
        {"$lookup": {"from": "borrowers", "localField": "borrower_id", "foreignField": "_id", "as": "borrower"}},
        {"$lookup": {"from": "tools", "localField": "tool_id", "foreignField": "_id", "as": "tool_doc"}},
        {"$unwind": {"path": "$borrower", "preserveNullAndEmptyArrays": True}},
        {"$unwind": {"path": "$tool_doc", "preserveNullAndEmptyArrays": True}},
        {"$project": {
            "borrower_name": "$borrower.borrower_name",
            "borrower_id": "$borrower.borrower_id",
            "tool_name": "$tool_doc.tool_name",
            "tool_code": "$tool_doc.tool_code",
            "barcode": 1,
            "borrow_date": 1,
            "expected_return_date": 1,
            "is_overdue": {
                "$cond": {
                    "if": {"$and": [
                        {"$ne": ["$expected_return_date", ""]},
                        {"$ne": ["$expected_return_date", None]},
                        {"$lt": ["$expected_return_date", date.today().isoformat()]},
                    ]},
                    "then": 1,
                    "else": 0,
                }
            },
        }},
    ]
    currently_borrowed = to_docs(db.transactions.aggregate(currently_borrowed_pipeline))

    borrow_summary = list(db.transactions.aggregate([
        {"$group": {"_id": {"$substrCP": ["$borrow_date", 0, 10]}, "count": {"$sum": 1}}},
        {"$sort": {"_id": ASCENDING}},
        {"$limit": 30},
        {"$project": {"date": "$_id", "count": 1, "_id": 0}},
    ]))

    return_summary = list(db.transactions.aggregate([
        {"$match": {"return_date": {"$ne": None}}},
        {"$group": {"_id": {"$substrCP": ["$return_date", 0, 10]}, "count": {"$sum": 1}}},
        {"$sort": {"_id": ASCENDING}},
        {"$limit": 30},
        {"$project": {"date": "$_id", "count": 1, "_id": 0}},
    ]))

    return render_template(
        "reports.html",
        most_borrowed_tools=most_borrowed_tools,
        currently_borrowed=currently_borrowed,
        borrow_summary=borrow_summary,
        return_summary=return_summary,
    )


@app.route("/export/tools.csv")
@login_required
def export_tools_csv():
    db = get_db()
    tools_list = to_docs(db.tools.find({}).sort("_id", DESCENDING))

    view_mode = request.args.get("view", "csv").strip().lower()
    if view_mode == "table":
        headers = ["Tool Name", "Tool Code", "Category", "Description", "Quantity", "Available Quantity", "Barcode", "Barcode Image", "Status", "Date Added"]
        table_rows = [[
            row["tool_name"], row["tool_code"], row["category"], row.get("description"),
            row["quantity"], row["available_quantity"], row["barcode"],
            row.get("barcode_image"), row["status"], row.get("date_added"),
        ] for row in tools_list]
        return render_template(
            "export_table.html",
            title="Tools Export (Printable Table)",
            filename="tools_export.csv",
            headers=headers,
            rows=table_rows,
            generated_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            csv_download_url=url_for("export_tools_csv"),
        )

    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(["Tool Name", "Tool Code", "Category", "Description", "Quantity", "Available Quantity", "Barcode", "Barcode Image", "Status", "Date Added"])
    for row in tools_list:
        writer.writerow([
            row["tool_name"], row["tool_code"], row["category"], row.get("description"),
            row["quantity"], row["available_quantity"], row["barcode"],
            row.get("barcode_image"), row["status"], row.get("date_added"),
        ])
    csv_content = output.getvalue()
    output.close()
    return Response(csv_content, mimetype="text/csv", headers={"Content-Disposition": "attachment; filename=tools_export.csv"})


@app.route("/export/transactions.csv")
@login_required
def export_transactions_csv():
    db = get_db()
    pipeline = [*_transaction_join_stages(), {"$sort": {"_id": DESCENDING}}]
    rows = to_docs(db.transactions.aggregate(pipeline))

    view_mode = request.args.get("view", "csv").strip().lower()
    if view_mode == "table":
        headers = ["Transaction ID", "Borrower Name", "Borrower ID", "Tool Name", "Barcode", "Borrow Date", "Expected Return Date", "Return Date", "Status", "Lent By"]
        table_rows = [[
            row["id"], row.get("borrower_name"), row.get("borrower_id"),
            row.get("tool_name"), row.get("barcode"), row.get("borrow_date"),
            row.get("expected_return_date"), row.get("return_date"),
            row.get("status"), row.get("lent_by_admin_name"),
        ] for row in rows]
        return render_template(
            "export_table.html",
            title="Transactions Export (Printable Table)",
            filename="transactions_export.csv",
            headers=headers,
            rows=table_rows,
            generated_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            csv_download_url=url_for("export_transactions_csv"),
        )

    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(["Transaction ID", "Borrower Name", "Borrower ID", "Tool Name", "Barcode", "Borrow Date", "Expected Return Date", "Return Date", "Status", "Lent By"])
    for row in rows:
        writer.writerow([
            row["id"], row.get("borrower_name"), row.get("borrower_id"),
            row.get("tool_name"), row.get("barcode"), row.get("borrow_date"),
            row.get("expected_return_date"), row.get("return_date"),
            row.get("status"), row.get("lent_by_admin_name"),
        ])
    csv_content = output.getvalue()
    output.close()
    return Response(csv_content, mimetype="text/csv", headers={"Content-Disposition": "attachment; filename=transactions_export.csv"})


@app.route("/api/tool/<barcode_val>")
@login_required
def api_tool_by_barcode(barcode_val):
    db = get_db()
    tool = get_tool_by_barcode(db, barcode_val)
    if tool is None:
        return {"found": False}
    return {"found": True, "tool": build_tool_payload(tool)}


@app.route("/api/borrowed/<barcode_val>")
@login_required
def api_borrowed_by_barcode(barcode_val):
    db = get_db()
    row = get_active_borrow_transaction(db, barcode_val)
    if row is None:
        return {"found": False}

    is_overdue = False
    expected_date = parse_date_input(row.get("expected_return_date"))
    if expected_date and date.today() > expected_date:
        is_overdue = True

    return {
        "found": True,
        "transaction": {
            "id": row["id"],
            "borrow_date": row.get("borrow_date"),
            "expected_return_date": row.get("expected_return_date"),
            "status": row.get("status"),
            "is_overdue": is_overdue,
            "borrower_name": row.get("borrower_name"),
            "borrower_id": row.get("borrower_id"),
            "course_department": row.get("course_department"),
            "contact_number": row.get("contact_number"),
            "tool_name": row.get("tool_name"),
            "tool_code": row.get("tool_code"),
            "category": row.get("category"),
            "barcode": row.get("barcode"),
            "lent_by_admin_name": row.get("lent_by_admin_name"),
        },
    }


@app.route("/api/borrower/<borrower_id>")
@login_required
def api_borrower_by_id(borrower_id):
    db = get_db()
    borrower = db.borrowers.find_one({"borrower_id": borrower_id.strip()})
    
    if borrower is None:
        return {"found": False}, 404
    
    return {
        "found": True,
        "borrower_id": borrower.get("borrower_id"),
        "borrower_name": borrower.get("borrower_name"),
        "course_department": borrower.get("course_department"),
        "contact_number": borrower.get("contact_number"),
    }


@app.template_filter("status_label")
def status_label(status_value):
    status_map = {
        "borrowed": "Borrowed",
        "returned": "Returned",
        "returned_overdue": "Returned (Overdue)",
    }
    return status_map.get(status_value, str(status_value).title())


@app.template_filter("status_badge")
def status_badge(status_value):
    if status_value == "borrowed":
        return "bg-amber-100 text-amber-800"
    if status_value == "returned_overdue":
        return "bg-red-100 text-red-800"
    if status_value == "returned":
        return "bg-emerald-100 text-emerald-800"
    return "bg-slate-100 text-slate-800"


if __name__ == "__main__":
    app.run(
        host=os.environ.get("FLASK_HOST", "127.0.0.1"),
        port=int(os.environ.get("PORT", "5000")),
        debug=os.environ.get("FLASK_DEBUG", "1") == "1",
    )
